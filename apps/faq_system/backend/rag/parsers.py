"""FAQ システムファイルパーサー.

Microsoft markitdown を主パーサーとして使用し、
PDF/Word/Excel/PPT/CSV/JSON/HTML 等を Markdown に変換する。
markitdown 未対応 or 失敗時のみレガシーパーサーにフォールバック。

統合インターフェース ``ParseResult`` で全形式の結果を返す。
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import IO, Any


logger = logging.getLogger(__name__)

# --- markitdown (Microsoft) ---
try:
    from markitdown import MarkItDown as _MarkItDown

    _markitdown_instance: _MarkItDown | None = _MarkItDown()
except ImportError:
    _markitdown_instance = None
    logger.info("markitdown 未インストール — レガシーパーサーを使用")

# --- レガシー依存（フォールバック用） ---
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import openpyxl
except ImportError:
    openpyxl = None


# ---------------------------------------------------------------------------
# 統合パース結果
# ---------------------------------------------------------------------------


@dataclass
class ParseResult:
    """パーサー出力の統一型.

    Attributes:
        content: 抽出されたテキスト (Markdown 形式)
        metadata: ファイル由来のメタ情報（タイトル・著者・ページ数等）
        file_type: MIME タイプまたは拡張子
        sections: 構造的セクション（見出し・シート名など）
    """

    content: str
    metadata: dict[str, Any] = field(default_factory=dict)
    file_type: str = ""
    sections: list[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# markitdown ベースパーサー
# ---------------------------------------------------------------------------


def _parse_with_markitdown(
    file_bytes: bytes,
    filename: str,
) -> ParseResult | None:
    """markitdown で変換を試みる.

    成功時は ParseResult、失敗時は None を返す。
    """
    if _markitdown_instance is None:
        return None

    ext = Path(filename).suffix.lower()

    try:
        # markitdown はファイルパスまたはストリームを受け付ける
        # ストリームの場合は StreamInfo で拡張子ヒントを渡す
        from markitdown import StreamInfo

        stream = io.BytesIO(file_bytes)
        stream_info = StreamInfo(
            extension=ext,
            filename=filename,
        )
        result = _markitdown_instance.convert_stream(stream, stream_info=stream_info)

        markdown_text = result.markdown or ""
        if not markdown_text.strip():
            return None

        metadata: dict[str, Any] = {"parser": "markitdown"}
        if result.title:
            metadata["title"] = result.title

        # Markdown からセクション見出しを抽出
        sections: list[str] = []
        for line in markdown_text.split("\n"):
            stripped = line.strip()
            if stripped.startswith("#"):
                heading = stripped.lstrip("#").strip()
                if heading:
                    sections.append(heading)

        return ParseResult(
            content=markdown_text,
            metadata=metadata,
            file_type=_ext_to_mime(ext),
            sections=sections,
        )
    except Exception:
        logger.debug("markitdown 変換失敗 (%s) — フォールバック", filename, exc_info=True)
        return None


def _parse_with_markitdown_path(
    file_path: str | Path,
) -> ParseResult | None:
    """markitdown でローカルファイルパスから変換."""
    if _markitdown_instance is None:
        return None

    p = Path(file_path)
    try:
        result = _markitdown_instance.convert(str(p))
        markdown_text = result.markdown or ""
        if not markdown_text.strip():
            return None

        metadata: dict[str, Any] = {"parser": "markitdown"}
        if result.title:
            metadata["title"] = result.title

        sections: list[str] = []
        for line in markdown_text.split("\n"):
            stripped = line.strip()
            if stripped.startswith("#"):
                heading = stripped.lstrip("#").strip()
                if heading:
                    sections.append(heading)

        return ParseResult(
            content=markdown_text,
            metadata=metadata,
            file_type=_ext_to_mime(p.suffix.lower()),
            sections=sections,
        )
    except Exception:
        logger.debug("markitdown パス変換失敗 (%s) — フォールバック", file_path, exc_info=True)
        return None


# ---------------------------------------------------------------------------
# レガシーパーサー（フォールバック用）
# ---------------------------------------------------------------------------


class _LegacyParser:
    """markitdown が使えない場合のフォールバックパーサー."""

    @staticmethod
    def parse_pdf(file_stream: IO[bytes]) -> ParseResult:
        """PDF テキスト抽出."""
        metadata: dict[str, Any] = {"parser": "pdf"}

        if pdfplumber is not None:
            try:
                file_stream.seek(0)
                with pdfplumber.open(file_stream) as pdf:
                    pages: list[str] = []
                    sections: list[str] = []
                    for i, page in enumerate(pdf.pages):
                        text = page.extract_text() or ""
                        if text:
                            pages.append(text)
                            sections.append(f"Page {i + 1}")
                    metadata["page_count"] = len(pdf.pages)
                    pdf_meta = getattr(pdf, "metadata", None)
                    if pdf_meta:
                        _merge_pdf_metadata(pdf_meta, metadata)
                    return ParseResult(
                        content="\n\n".join(pages),
                        metadata=metadata,
                        file_type="application/pdf",
                        sections=sections,
                    )
            except Exception:
                logger.debug("pdfplumber 失敗、pypdf にフォールバック", exc_info=True)

        if PdfReader is None:
            msg = "PDF パースには markitdown または pdfplumber / pypdf が必要"
            raise ImportError(msg)

        file_stream.seek(0)
        reader = PdfReader(file_stream)
        pages_legacy: list[str] = []
        sections_legacy: list[str] = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text:
                pages_legacy.append(text)
                sections_legacy.append(f"Page {i + 1}")

        metadata["page_count"] = len(reader.pages)
        reader_metadata = reader.metadata
        if reader_metadata is not None:
            _merge_pdf_metadata(
                {key: getattr(reader_metadata, key, None) for key in ("title", "author", "subject", "creator")},
                metadata,
            )

        return ParseResult(
            content="\n\n".join(pages_legacy),
            metadata=metadata,
            file_type="application/pdf",
            sections=sections_legacy,
        )

    @staticmethod
    def parse_docx(file_stream: IO[bytes]) -> ParseResult:
        """Word テキスト抽出."""
        if DocxDocument is None:
            msg = "DOCX パースには markitdown または python-docx が必要"
            raise ImportError(msg)

        file_stream.seek(0)
        doc = DocxDocument(file_stream)
        paragraphs: list[str] = []
        sections: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            paragraphs.append(text)
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                sections.append(text)

        for table in doc.tables:
            rows: list[str] = []
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                paragraphs.append("\n".join(rows))

        metadata: dict[str, Any] = {"parser": "python-docx"}
        props = doc.core_properties
        if props and props.title:
            metadata["title"] = props.title
        if props and props.author:
            metadata["author"] = props.author

        return ParseResult(
            content="\n".join(paragraphs),
            metadata=metadata,
            file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            sections=sections,
        )

    @staticmethod
    def parse_excel(file_stream: IO[bytes]) -> ParseResult:
        """Excel テキスト抽出."""
        if openpyxl is None:
            msg = "Excel パースには markitdown または openpyxl が必要"
            raise ImportError(msg)

        file_stream.seek(0)
        wb = openpyxl.load_workbook(file_stream, read_only=True, data_only=True)
        metadata: dict[str, Any] = {
            "parser": "openpyxl",
            "sheet_count": len(wb.sheetnames),
            "sheet_names": list(wb.sheetnames),
        }
        sections: list[str] = []
        sheet_texts: list[str] = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sections.append(sheet_name)
            rows_data: list[list[str]] = []
            for row in ws.iter_rows(values_only=True):
                vals = [str(c) if c is not None else "" for c in row]
                if any(v.strip() for v in vals):
                    rows_data.append(vals)
            if not rows_data:
                continue
            headers = rows_data[0]
            row_texts: list[str] = []
            for data_row in rows_data[1:]:
                parts: list[str] = []
                for i, val in enumerate(data_row):
                    if not val.strip():
                        continue
                    h = headers[i] if i < len(headers) and headers[i].strip() else f"Col{i + 1}"
                    parts.append(f"{h}: {val}")
                if parts:
                    row_texts.append(", ".join(parts))
            if row_texts:
                sheet_texts.append(f"[シート: {sheet_name}]\n" + "\n".join(row_texts))
        wb.close()

        metadata["total_rows"] = sum(1 for st in sheet_texts for _line in st.split("\n") if _line.strip())
        return ParseResult(
            content="\n\n".join(sheet_texts),
            metadata=metadata,
            file_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            sections=sections,
        )


# ---------------------------------------------------------------------------
# 統合パーサー API
# ---------------------------------------------------------------------------


class FileParser:
    """ファイルパーサー統合クラス.

    markitdown を優先し、失敗時のみレガシーパーサーにフォールバック。
    """

    # ---- 後方互換 API（文字列返却） ----------------------------------------

    @staticmethod
    def parse_pdf(file_stream: IO[bytes]) -> str:
        """PDF テキスト抽出（後方互換）."""
        return FileParser.parse_pdf_rich(file_stream).content

    @staticmethod
    def parse_docx(file_stream: IO[bytes]) -> str:
        """Word テキスト抽出（後方互換）."""
        return FileParser.parse_docx_rich(file_stream).content

    @staticmethod
    def parse_excel(file_stream: IO[bytes]) -> str:
        """Excel テキスト抽出（後方互換）."""
        return FileParser.parse_excel_rich(file_stream).content

    @staticmethod
    def parse_csv(file_stream: IO[str]) -> str:
        """CSV テキスト抽出（後方互換）."""
        import csv

        reader = csv.reader(file_stream)
        rows = list(reader)
        if not rows:
            return ""
        headers = rows[0]
        chunks: list[str] = []
        for row in rows[1:]:
            parts = [f"{headers[i]}: {v}" if i < len(headers) else v for i, v in enumerate(row)]
            chunks.append(", ".join(parts))
        return "\n---\n".join(chunks)

    @staticmethod
    def parse_text(file_content: str) -> str:
        """プレーンテキスト（後方互換）."""
        return file_content

    # ---- Rich API（ParseResult 返却）--------------------------------------

    @staticmethod
    def parse_pdf_rich(file_stream: IO[bytes]) -> ParseResult:
        """PDF → Markdown 変換."""
        file_stream.seek(0)
        data = file_stream.read()
        result = _parse_with_markitdown(data, "doc.pdf")
        if result is not None:
            return result
        return _LegacyParser.parse_pdf(io.BytesIO(data))

    @staticmethod
    def parse_docx_rich(file_stream: IO[bytes]) -> ParseResult:
        """Word → Markdown 変換."""
        file_stream.seek(0)
        data = file_stream.read()
        result = _parse_with_markitdown(data, "doc.docx")
        if result is not None:
            return result
        return _LegacyParser.parse_docx(io.BytesIO(data))

    @staticmethod
    def parse_excel_rich(file_stream: IO[bytes]) -> ParseResult:
        """Excel → Markdown 変換."""
        file_stream.seek(0)
        data = file_stream.read()
        result = _parse_with_markitdown(data, "doc.xlsx")
        if result is not None:
            return result
        return _LegacyParser.parse_excel(io.BytesIO(data))

    @staticmethod
    def parse_text_rich(file_content: str, filename: str = "") -> ParseResult:
        """プレーンテキスト + メタデータ."""
        return ParseResult(
            content=file_content,
            metadata={"parser": "text", "char_count": len(file_content)},
            file_type="text/plain",
        )

    # ---- 自動判定 ----------------------------------------------------------

    @staticmethod
    def parse_auto(
        file_bytes: bytes,
        filename: str,
    ) -> ParseResult:
        """ファイル名の拡張子から適切なパーサーを自動選択.

        markitdown を優先し、失敗時にレガシーパーサーへフォールバック。

        Args:
            file_bytes: ファイルバイナリ
            filename: ファイル名（拡張子判定用）

        Returns:
            ParseResult

        Raises:
            ValueError: テキスト抽出に完全失敗した場合
        """
        # markitdown を最優先で試行
        result = _parse_with_markitdown(file_bytes, filename)
        if result is not None:
            return result

        # レガシーフォールバック
        ext = Path(filename).suffix.lower()
        try:
            if ext == ".pdf":
                return _LegacyParser.parse_pdf(io.BytesIO(file_bytes))
            if ext in {".docx", ".doc"}:
                return _LegacyParser.parse_docx(io.BytesIO(file_bytes))
            if ext in {".xlsx", ".xls"}:
                return _LegacyParser.parse_excel(io.BytesIO(file_bytes))
            if ext == ".csv":
                import csv as _csv

                text = file_bytes.decode("utf-8")
                reader = _csv.reader(io.StringIO(text))
                rows = list(reader)
                if not rows:
                    return ParseResult(content="", metadata={"parser": "csv"}, file_type="text/csv")
                headers = rows[0]
                chunks: list[str] = []
                for row in rows[1:]:
                    parts = [f"{headers[i]}: {v}" if i < len(headers) else v for i, v in enumerate(row)]
                    chunks.append(", ".join(parts))
                return ParseResult(
                    content="\n---\n".join(chunks),
                    metadata={"parser": "csv", "row_count": len(rows) - 1, "columns": headers},
                    file_type="text/csv",
                    sections=headers,
                )
            # テキスト系
            text = file_bytes.decode("utf-8")
            return ParseResult(
                content=text,
                metadata={"parser": "text"},
                file_type="text/plain",
            )
        except ImportError:
            raise
        except UnicodeDecodeError as exc:
            msg = f"ファイル '{filename}' のテキスト抽出に失敗（エンコーディング不正）"
            raise ValueError(msg) from exc
        except Exception as exc:
            msg = f"ファイル '{filename}' のパースに失敗: {exc}"
            raise ValueError(msg) from exc

    # ---- パス指定の直接変換 ------------------------------------------------

    @staticmethod
    def parse_file(file_path: str | Path) -> ParseResult:
        """ファイルパスから直接変換.

        markitdown を使う場合はテンポラリファイル不要。
        """
        p = Path(file_path)
        # markitdown はパスを直接受け付けるので最高効率
        result = _parse_with_markitdown_path(p)
        if result is not None:
            return result
        # フォールバック: バイトを読んで parse_auto
        return FileParser.parse_auto(p.read_bytes(), p.name)


# ---------------------------------------------------------------------------
# ヘルパー
# ---------------------------------------------------------------------------


_MIME_MAP: dict[str, str] = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".doc": "application/msword",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".xls": "application/vnd.ms-excel",
    ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    ".csv": "text/csv",
    ".json": "application/json",
    ".html": "text/html",
    ".htm": "text/html",
    ".md": "text/markdown",
    ".txt": "text/plain",
}


def _ext_to_mime(ext: str) -> str:
    """拡張子から MIME タイプに変換."""
    return _MIME_MAP.get(ext, "application/octet-stream")


def _merge_pdf_metadata(source: dict[str, Any] | Any, target: dict[str, Any]) -> None:
    """PDF メタデータを target 辞書へ安全にマージする."""
    if not isinstance(source, dict):
        return

    for key in ("title", "author", "subject", "creator"):
        value = source.get(key)
        if isinstance(value, str) and value.strip():
            target[key] = value.strip()
