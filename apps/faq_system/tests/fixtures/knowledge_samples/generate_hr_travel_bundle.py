"""人事・出張精算サンプル文書を生成する."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from openpyxl import Workbook
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas


def generate_pdf(output_path: Path) -> None:
    """正式規程の PDF を生成する."""
    lines = [
        "Corporate Travel Expense Policy 2025 Official Version",
        "Effective date: 2025-01-01",
        "",
        "Section 2.1 Meal allowance",
        "- Lunch allowance: 1500 JPY per travel day.",
        "- Dinner allowance: 3000 JPY per overnight travel day.",
        "- No meal allowance can be claimed when a host or conference already provides the meal.",
        "",
        "Section 3 Required receipts",
        "- Lodging, taxi, airfare, and rail tickets require an original or approved digital receipt.",
        "- Missing lodging receipts require a payment statement and manager approval.",
        "",
        "Section 4 Taxi use",
        "- Taxi reimbursement is allowed only after 22:00 or when public transit is unavailable.",
        "- Taxi expenses above 8000 JPY require manager approval.",
        "",
        "Section 5 Lodging caps",
        "- Use the official allowance matrix by role, city, and overnight stay.",
        "- When a later update notice conflicts with this policy, the latest effective notice wins.",
    ]

    pdf = canvas.Canvas(str(output_path), pagesize=A4)
    _, height = A4
    text = pdf.beginText(18 * mm, height - 20 * mm)
    text.setFont("Helvetica", 11)
    for line in lines:
        text.textLine(line)
    pdf.drawText(text)
    pdf.showPage()
    pdf.save()


def generate_docx(output_path: Path) -> None:
    """FAQ と例外ケースの DOCX を生成する."""
    document = Document()
    document.add_heading("Travel Expense FAQ and Exception Notes", level=1)
    document.add_paragraph("Effective usage notes for HR and travel operations.")

    faq_items = [
        (
            "Q1. What if the hotel receipt is missing?",
            "Submit the hotel payment statement within 3 business days and attach manager approval before reimbursement.",
        ),
        (
            "Q2. Can a same-day Osaka trip use a taxi home?",
            "Only when the employee arrives after 22:00 or carries equipment that cannot be taken on public transit.",
        ),
        (
            "Q3. If a conference includes lunch, can lunch allowance still be claimed?",
            "No. Duplicate meal claims are not allowed when lunch is already provided.",
        ),
        (
            "Q4. Which rule wins when the update notice conflicts with the official policy?",
            "Use the latest update notice by effective date and keep the notice together with the policy in the same review batch.",
        ),
    ]
    for question, answer in faq_items:
        document.add_heading(question, level=2)
        document.add_paragraph(answer)

    table = document.add_table(rows=1, cols=3)
    headers = table.rows[0].cells
    headers[0].text = "Case"
    headers[1].text = "Extra evidence"
    headers[2].text = "Decision"

    rows = [
        ("Missing hotel receipt", "Payment statement + manager approval", "Allowed after review"),
        ("Taxi before 22:00", "Transit outage proof", "Only if public transit unavailable"),
        ("Conference lunch included", "Conference agenda", "Lunch allowance rejected"),
    ]
    for case_name, evidence, decision in rows:
        row = table.add_row().cells
        row[0].text = case_name
        row[1].text = evidence
        row[2].text = decision

    document.save(output_path)


def generate_xlsx(output_path: Path) -> None:
    """上限表の Excel を生成する."""
    workbook = Workbook()
    lodging_sheet = workbook.active
    lodging_sheet.title = "LodgingCaps"
    lodging_sheet.append(["Role", "City", "Overnight", "MaxJPY", "PolicyVersion"])
    for row in [
        ["Employee", "Tokyo", "Yes", 15000, "baseline_2025"],
        ["Manager", "Tokyo", "Yes", 17000, "baseline_2025"],
        ["Director", "Shanghai", "Yes", 22000, "baseline_2025"],
        ["Manager", "Osaka", "No", 0, "baseline_2025"],
    ]:
        lodging_sheet.append(row)

    allowance_sheet = workbook.create_sheet("DailyAllowance")
    allowance_sheet.append(["Role", "DomesticPerDayJPY", "InternationalPerDayJPY", "Notes"])
    for row in [
        ["Employee", 2500, 5000, "Meal allowance requires no host-provided meal"],
        ["Manager", 3000, 6500, "Manager approval required above standard cap"],
        ["Director", 4000, 8000, "Use latest notice if updated after baseline"],
    ]:
        allowance_sheet.append(row)

    workbook.save(output_path)


def generate_txt(output_path: Path) -> None:
    """更新通知テキストを生成する."""
    content = "\n".join(
        [
            "Travel Policy Update Notice 2026-04",
            "Effective date: 2026-04-01",
            "",
            "1. Manager Tokyo overnight lodging cap changes from 17000 JPY to 19000 JPY.",
            "2. Digital lodging receipts are accepted, but a hotel payment statement is still required when the invoice is missing.",
            "3. Taxi reimbursement after 22:00 remains allowed under the official policy conditions.",
            "4. This notice overrides older baseline_2025 matrix values for the affected cases only.",
        ]
    )
    output_path.write_text(content, encoding="utf-8")


def main() -> None:
    """サンプル文書セットを生成する."""
    output_dir = Path(__file__).resolve().parent / "hr_travel_bundle"
    output_dir.mkdir(parents=True, exist_ok=True)

    generate_pdf(output_dir / "travel_policy_official_2025.pdf")
    generate_docx(output_dir / "travel_faq_exceptions.docx")
    generate_xlsx(output_dir / "travel_allowance_matrix.xlsx")
    generate_txt(output_dir / "travel_policy_update_notice_2026.txt")


if __name__ == "__main__":
    main()
