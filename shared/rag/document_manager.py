"""ドキュメントライフサイクル管理.

アップロード → パース → チャンク → プレビュー → インデックス → 再インデックス
のフルサイクルを管理する。

ファイルストレージ:
    アップロードされたファイルは data/rag/{app_name}/{collection}/uploads/ に保存される。
    storage_root を指定することでベースディレクトリをカスタマイズ可能。

使用例:
    >>> from shared.rag.document_manager import DocumentManager
    >>>
    >>> doc_mgr = DocumentManager(collection_manager=col_mgr, session_factory=sf)
    >>> doc = await doc_mgr.upload_document(
    ...     collection_name="internal_kb",
    ...     file_content=b"...",
    ...     filename="guide.pdf",
    ...     user_id="admin",
    ... )
    >>> chunks = await doc_mgr.preview_chunks(doc.document_id)
"""

from __future__ import annotations

import hashlib
import json
import logging
import mimetypes
import uuid
from datetime import UTC, datetime
from pathlib import Path
from typing import TYPE_CHECKING, Any

from sqlalchemy import select

from shared.rag.models import (
    DocumentRecordModel,
    DocumentStatus,
)


if TYPE_CHECKING:
    from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker

    from shared.rag.collection_manager import CollectionManager

logger = logging.getLogger(__name__)

# デフォルトストレージルート
_DEFAULT_STORAGE_ROOT = Path("data/rag")


class DocumentManager:
    """ドキュメントライフサイクル管理.

    upload → parse → chunk → preview → index → query → re-index の
    フルサイクルを CollectionManager と連携して管理する。

    Args:
        collection_manager: CollectionManager インスタンス
        session_factory: SQLAlchemy async セッションファクトリ
    """

    def __init__(
        self,
        collection_manager: CollectionManager,
        session_factory: async_sessionmaker[AsyncSession],
        storage_root: Path | str | None = None,
    ) -> None:
        self._col_mgr = collection_manager
        self._session_factory = session_factory
        self._storage_root = Path(storage_root) if storage_root else _DEFAULT_STORAGE_ROOT

    # ------------------------------------------------------------------
    # アップロード
    # ------------------------------------------------------------------

    async def upload_document(
        self,
        *,
        collection_name: str,
        file_content: bytes,
        filename: str,
        metadata: dict[str, Any] | None = None,
        user_id: str | None = None,
    ) -> DocumentRecordModel:
        """ドキュメントをアップロード.

        Args:
            collection_name: 対象コレクション名
            file_content: ファイルバイナリ
            filename: ファイル名
            metadata: メタデータ
            user_id: アップロードユーザーID

        Returns:
            作成された DocumentRecordModel

        Raises:
            ValueError: コレクションが存在しない場合
        """
        # コレクション存在確認
        collection = await self._col_mgr.get_collection(collection_name)
        if collection is None:
            msg = f"Collection '{collection_name}' not found"
            raise ValueError(msg)

        # ファイルタイプ判定
        mime_type, _ = mimetypes.guess_type(filename)
        file_type = mime_type or "application/octet-stream"

        # SHA-256 ハッシュ（重複検出用）
        content_hash = hashlib.sha256(file_content).hexdigest()
        normalized_metadata = dict(metadata or {})
        document_group_id = self._extract_document_group_id(normalized_metadata)
        tags = self._extract_tags(normalized_metadata)

        doc_id = uuid.uuid4().hex[:16]

        # ファイルをディスクに保存
        app_name = collection.app_name or "default"
        upload_dir = self._storage_root / app_name / collection_name / "uploads"
        upload_dir.mkdir(parents=True, exist_ok=True)
        file_path = upload_dir / f"{doc_id}_{filename}"
        file_path.write_bytes(file_content)

        async with self._session_factory() as session:
            record = DocumentRecordModel(
                document_id=doc_id,
                collection_name=collection_name,
                filename=filename,
                file_type=file_type,
                file_size=len(file_content),
                status=DocumentStatus.UPLOADED.value,
                content_hash=content_hash,
                document_group_id=document_group_id,
                tags_json=json.dumps(tags, ensure_ascii=False),
                metadata_json=json.dumps(
                    {**normalized_metadata, "file_path": str(file_path)},
                    ensure_ascii=False,
                ),
                uploaded_by=user_id,
            )
            session.add(record)
            await session.commit()
            await session.refresh(record)
            return record

    # ------------------------------------------------------------------
    # チャンクプレビュー
    # ------------------------------------------------------------------

    async def preview_chunks(
        self,
        document_id: str,
        chunk_strategy: str | None = None,
        chunk_size: int | None = None,
        chunk_overlap: int | None = None,
    ) -> list[dict[str, Any]]:
        """ドキュメントのチャンクプレビューを生成.

        コレクションのデフォルト設定またはカスタム設定でチャンキングを試行し、
        インデックス前に結果を確認できる。

        Args:
            document_id: ドキュメントID
            chunk_strategy: チャンキング戦略（None でコレクション設定を使用）
            chunk_size: チャンクサイズ（None でコレクション設定を使用）
            chunk_overlap: オーバーラップ（None でコレクション設定を使用）

        Returns:
            チャンクプレビューのリスト

        Raises:
            ValueError: ドキュメントが存在しない場合
        """
        doc = await self.get_document(document_id)
        if doc is None:
            msg = f"Document '{document_id}' not found"
            raise ValueError(msg)

        # コレクション設定を取得
        collection = await self._col_mgr.get_collection(doc.collection_name)
        effective_strategy = chunk_strategy or (collection.chunk_strategy if collection else "recursive")
        effective_size = chunk_size or (collection.chunk_size if collection else 1000)
        effective_overlap = chunk_overlap or (collection.chunk_overlap if collection else 200)

        # ファイルストレージからコンテンツを読み込む
        content = self._read_document_content(doc)

        # 簡易チャンキング（プレビュー用）
        chunks = self._split_text(
            text=content,
            chunk_size=effective_size,
            chunk_overlap=effective_overlap,
        )

        return [
            {
                "index": i,
                "content": chunk,
                "length": len(chunk),
                "strategy": effective_strategy,
            }
            for i, chunk in enumerate(chunks)
        ]

    # ------------------------------------------------------------------
    # インデックス
    # ------------------------------------------------------------------

    async def index_document(
        self,
        document_id: str,
        chunks: list[dict[str, Any]] | None = None,
    ) -> DocumentRecordModel:
        """ドキュメントをインデックスに登録.

        RAGService を使用して実際にチャンク化 → embedding → vector DB 保存を行う。

        Args:
            document_id: ドキュメントID
            chunks: カスタムチャンク（None で自動チャンキング）

        Returns:
            更新された DocumentRecordModel

        Raises:
            ValueError: ドキュメントが存在しない場合
        """
        doc = await self.get_document(document_id)
        if doc is None:
            msg = f"Document '{document_id}' not found"
            raise ValueError(msg)

        # コレクション設定から RAGConfig を構築
        rag_config = await self._col_mgr.build_rag_config(doc.collection_name)

        # ファイルコンテンツを読み込み
        content = self._read_document_content(doc)
        if not content or content.startswith("[Binary:"):
            msg = f"Document '{document_id}' has no readable text content"
            raise ValueError(msg)

        # RAGService で実際に embedding + vector DB 保存
        from shared.services.rag_service import RAGService

        service = RAGService(rag_config)
        result = await service.execute(
            action="add_document",
            content=content,
            source=doc.filename,
            metadata=self._build_index_metadata(doc),
        )

        if not result.success:
            # インデックス失敗 → エラーステータスに更新
            async with self._session_factory() as session:
                db_result = await session.execute(
                    select(DocumentRecordModel).where(DocumentRecordModel.document_id == document_id)
                )
                db_doc = db_result.scalar_one_or_none()
                if db_doc:
                    db_doc.status = DocumentStatus.ERROR.value
                    db_doc.error_message = result.error_message or "Index failed"
                    await session.commit()
            msg = f"Index failed: {result.error_message}"
            raise ValueError(msg)

        # 成功 → INDEXED ステータスに更新
        chunk_ids = result.data.get("ids", []) if result.data else []
        chunk_count = result.data.get("count", 0) if result.data else 0

        async with self._session_factory() as session:
            db_result = await session.execute(
                select(DocumentRecordModel).where(DocumentRecordModel.document_id == document_id)
            )
            db_doc = db_result.scalar_one_or_none()
            if db_doc is None:
                msg = f"Document '{document_id}' not found"
                raise ValueError(msg)

            db_doc.status = DocumentStatus.INDEXED.value
            db_doc.chunk_count = chunk_count
            db_doc.chunk_ids_json = json.dumps(chunk_ids)
            db_doc.indexed_at = datetime.now(UTC)
            db_doc.error_message = None

            await session.commit()
            await session.refresh(db_doc)
            return db_doc

    async def reindex_document(self, document_id: str) -> DocumentRecordModel:
        """ドキュメントを再インデックス.

        既存のインデックスをリセットし、再度 index_document を実行する。

        Args:
            document_id: ドキュメントID

        Returns:
            更新された DocumentRecordModel

        Raises:
            ValueError: ドキュメントが存在しない場合
        """
        # まずステータスをリセット
        async with self._session_factory() as session:
            result = await session.execute(
                select(DocumentRecordModel).where(DocumentRecordModel.document_id == document_id)
            )
            doc = result.scalar_one_or_none()
            if doc is None:
                msg = f"Document '{document_id}' not found"
                raise ValueError(msg)

            doc.status = DocumentStatus.UPLOADED.value
            doc.chunk_count = 0
            doc.chunk_ids_json = "[]"
            doc.indexed_at = None

            await session.commit()

        # 再インデックス実行
        return await self.index_document(document_id)

    # ------------------------------------------------------------------
    # 一覧・取得・削除
    # ------------------------------------------------------------------

    async def list_documents(
        self,
        collection_name: str,
        status: str | None = None,
        limit: int = 100,
        offset: int = 0,
    ) -> list[DocumentRecordModel]:
        """ドキュメント一覧を取得.

        Args:
            collection_name: コレクション名
            status: ステータスでフィルタ
            limit: 取得件数
            offset: オフセット

        Returns:
            DocumentRecordModel のリスト
        """
        async with self._session_factory() as session:
            stmt = (
                select(DocumentRecordModel)
                .where(DocumentRecordModel.collection_name == collection_name)
                .offset(offset)
                .limit(limit)
            )
            if status is not None:
                stmt = stmt.where(DocumentRecordModel.status == status)
            result = await session.execute(stmt)
            return list(result.scalars().all())

    async def get_document(self, document_id: str) -> DocumentRecordModel | None:
        """ドキュメントを取得.

        Args:
            document_id: ドキュメントID

        Returns:
            DocumentRecordModel または None
        """
        async with self._session_factory() as session:
            result = await session.execute(
                select(DocumentRecordModel).where(DocumentRecordModel.document_id == document_id)
            )
            return result.scalar_one_or_none()

    async def delete_document(self, document_id: str) -> None:
        """ドキュメントを削除.

        DB レコードとディスク上のファイルの両方を削除する。

        Args:
            document_id: ドキュメントID

        Raises:
            ValueError: ドキュメントが存在しない場合
        """
        async with self._session_factory() as session:
            result = await session.execute(
                select(DocumentRecordModel).where(DocumentRecordModel.document_id == document_id)
            )
            doc = result.scalar_one_or_none()
            if doc is None:
                msg = f"Document '{document_id}' not found"
                raise ValueError(msg)

            # ディスク上のファイルを削除
            try:
                meta = json.loads(doc.metadata_json or "{}")
                file_path = meta.get("file_path")
                if file_path:
                    p = Path(file_path)
                    if p.exists():
                        p.unlink()
            except Exception:
                logger.warning("ファイル削除に失敗: doc_id=%s", document_id)

            await session.delete(doc)
            await session.commit()

    async def reindex_collection(self, collection_name: str) -> dict[str, Any]:
        """コレクション内の全ドキュメントを再インデックス.

        Args:
            collection_name: コレクション名

        Returns:
            再インデックス結果の辞書
        """
        docs = await self.list_documents(collection_name=collection_name, limit=10000)
        reindexed = 0
        errors = 0

        for doc in docs:
            try:
                await self.reindex_document(doc.document_id)
                reindexed += 1
            except Exception:
                logger.exception("ドキュメント '%s' の再インデックスに失敗", doc.document_id)
                errors += 1

        return {
            "collection_name": collection_name,
            "total": len(docs),
            "reindexed": reindexed,
            "errors": errors,
        }

    # ------------------------------------------------------------------
    # 内部ヘルパー
    # ------------------------------------------------------------------

    @staticmethod
    def _read_document_content(doc: DocumentRecordModel) -> str:
        """ファイルストレージからドキュメントコンテンツを読み込む.

        メタデータに保存された file_path からファイルを読み込む。
        バイナリファイル (PDF/DOCX/XLSX) はパーサーでテキスト抽出する。

        Args:
            doc: ドキュメントレコード

        Returns:
            ドキュメントテキスト（抽出失敗時は空文字列）
        """
        try:
            meta = json.loads(doc.metadata_json or "{}")
            file_path = meta.get("file_path")
            if not file_path:
                logger.warning("file_path が未設定: doc_id=%s", doc.document_id)
                return ""

            p = Path(file_path)
            if not p.exists():
                logger.warning("ファイルが存在しない: doc_id=%s path=%s", doc.document_id, file_path)
                return ""

            raw = p.read_bytes()

            # 拡張子で判定してパーサーを選択
            ext = p.suffix.lower()
            if ext in {".pdf", ".docx", ".doc", ".xlsx", ".xls"}:
                return _parse_binary_file(raw, doc.filename)

            # テキスト系: UTF-8 デコード
            try:
                return raw.decode("utf-8")
            except UnicodeDecodeError:
                # Shift_JIS フォールバック（日本語ファイル対策）
                try:
                    return raw.decode("shift_jis")
                except UnicodeDecodeError:
                    return raw.decode("utf-8", errors="replace")

        except Exception:
            logger.exception("ファイル読み込み失敗: doc_id=%s", doc.document_id)
            return ""

    @staticmethod
    def _split_text(
        text: str,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ) -> list[str]:
        """テキストを固定サイズでチャンク分割（プレビュー用簡易実装）.

        Args:
            text: 分割対象テキスト
            chunk_size: チャンクサイズ
            chunk_overlap: オーバーラップ

        Returns:
            チャンクのリスト
        """
        if not text:
            return []

        chunks: list[str] = []
        start = 0
        step = max(chunk_size - chunk_overlap, 1)

        while start < len(text):
            end = min(start + chunk_size, len(text))
            chunk = text[start:end]
            if chunk:
                chunks.append(chunk)
            start += step

        return chunks

    @staticmethod
    def _extract_document_group_id(metadata: dict[str, Any]) -> str | None:
        """metadata から document_group_id を取り出す."""
        value = metadata.get("document_group_id")
        if isinstance(value, str):
            stripped = value.strip()
            return stripped or None
        return None

    @staticmethod
    def _extract_tags(metadata: dict[str, Any]) -> list[str]:
        """metadata から tags を正規化して取り出す."""
        raw_tags = metadata.get("tags")
        if isinstance(raw_tags, list):
            return [tag.strip() for tag in raw_tags if isinstance(tag, str) and tag.strip()]
        if isinstance(raw_tags, str):
            return [tag.strip() for tag in raw_tags.split(",") if tag.strip()]
        return []

    @staticmethod
    def _build_index_metadata(doc: DocumentRecordModel) -> dict[str, Any]:
        """インデックス登録時に渡す metadata を構築する."""
        metadata = {
            "document_id": doc.document_id,
            "collection_name": doc.collection_name,
        }
        try:
            stored_metadata = json.loads(doc.metadata_json or "{}")
            if isinstance(stored_metadata, dict):
                metadata.update(stored_metadata)
        except (TypeError, ValueError):
            logger.warning("metadata_json の読込に失敗: doc_id=%s", doc.document_id)

        if doc.document_group_id:
            metadata["document_group_id"] = doc.document_group_id

        try:
            tags = json.loads(doc.tags_json or "[]")
            if isinstance(tags, list):
                metadata["tags"] = [tag for tag in tags if isinstance(tag, str) and tag.strip()]
        except (TypeError, ValueError):
            logger.warning("tags_json の読込に失敗: doc_id=%s", doc.document_id)

        return metadata


# ---------------------------------------------------------------------------
# モジュールレベルヘルパー
# ---------------------------------------------------------------------------


def _parse_binary_file(raw: bytes, filename: str) -> str:
    """バイナリファイル (PDF/DOCX/XLSX/PPTX 等) をテキストに変換.

    Microsoft markitdown を優先し、失敗時のみレガシーパーサーにフォールバック。

    Args:
        raw: ファイルバイナリ
        filename: ファイル名（拡張子判定用）

    Returns:
        抽出テキスト
    """
    import io as _io

    # markitdown を最優先で試行
    try:
        from markitdown import MarkItDown, StreamInfo

        md = MarkItDown()
        ext = Path(filename).suffix.lower()
        stream_info = StreamInfo(extension=ext, filename=filename)
        result = md.convert_stream(_io.BytesIO(raw), stream_info=stream_info)
        if result.markdown and result.markdown.strip():
            return result.markdown
    except ImportError:
        logger.debug("markitdown 未インストール — レガシーパーサーを使用")
    except Exception:
        logger.debug("markitdown 変換失敗 (%s) — フォールバック", filename, exc_info=True)

    # レガシーフォールバック
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        try:
            import pdfplumber

            with pdfplumber.open(_io.BytesIO(raw)) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
                return "\n\n".join(pages)
        except ImportError:
            pass
        try:
            from pypdf import PdfReader

            reader = PdfReader(_io.BytesIO(raw))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n\n".join(pages)
        except ImportError:
            logger.warning("PDF パースライブラリが未インストール: %s", filename)
            return ""

    if ext in {".docx", ".doc"}:
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(_io.BytesIO(raw))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        paragraphs.append(" | ".join(cells))
            return "\n".join(paragraphs)
        except ImportError:
            logger.warning("python-docx が未インストール: %s", filename)
            return ""

    if ext in {".xlsx", ".xls"}:
        try:
            import openpyxl as _openpyxl

            wb = _openpyxl.load_workbook(_io.BytesIO(raw), read_only=True, data_only=True)
            texts: list[str] = []
            for sn in wb.sheetnames:
                ws = wb[sn]
                rows_data: list[list[str]] = []
                for row in ws.iter_rows(values_only=True):
                    vals = [str(c) if c is not None else "" for c in row]
                    if any(v.strip() for v in vals):
                        rows_data.append(vals)
                if not rows_data:
                    continue
                headers = rows_data[0]
                row_strs: list[str] = []
                for dr in rows_data[1:]:
                    parts = []
                    for i, v in enumerate(dr):
                        if not v.strip():
                            continue
                        h = headers[i] if i < len(headers) and headers[i].strip() else f"Col{i + 1}"
                        parts.append(f"{h}: {v}")
                    if parts:
                        row_strs.append(", ".join(parts))
                if row_strs:
                    texts.append(f"[シート: {sn}]\n" + "\n".join(row_strs))
            wb.close()
            return "\n\n".join(texts)
        except ImportError:
            logger.warning("openpyxl が未インストール: %s", filename)
            return ""

    return ""
